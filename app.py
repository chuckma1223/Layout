import os
import re
import time
import tempfile
from pathlib import Path
from flask import Flask, render_template, request, send_file, send_from_directory, redirect, flash, url_for
from werkzeug.utils import secure_filename
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

from ai_service import parse_sections_with_ai

try:
    import docx2txt
except ImportError:
    docx2txt = None
import zipfile
import xml.etree.ElementTree as ET
try:
    import mammoth
except Exception:
    mammoth = None

ALLOWED_ARTICLE_EXTENSIONS = {"doc", "docx", "txt"}
ALLOWED_TEMPLATE_EXTENSIONS = {"doc", "docx"}
UPLOAD_FOLDER = Path(__file__).resolve().parent / "uploads"
FORMATTED_FOLDER = UPLOAD_FOLDER / "formatted_results"
BUILTIN_TEMPLATE_DIR = Path(__file__).resolve().parent / "template_files"

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024
app.secret_key = os.urandom(24)

UPLOAD_FOLDER.mkdir(exist_ok=True)
FORMATTED_FOLDER.mkdir(exist_ok=True)
BUILTIN_TEMPLATE_DIR.mkdir(exist_ok=True)

# cleanup stale Word lock / temporary files that may interfere with built-in template access
for path in BUILTIN_TEMPLATE_DIR.iterdir():
    if path.name.startswith('~$'):
        try:
            path.unlink()
        except Exception:
            pass
for path in UPLOAD_FOLDER.iterdir():
    if path.name.startswith('~$'):
        try:
            path.unlink()
        except Exception:
            pass


def get_builtin_templates():
    templates = []
    for path in BUILTIN_TEMPLATE_DIR.iterdir():
        if path.name.startswith('~$'):
            continue
        if path.suffix.lower() in {".doc", ".docx"}:
            templates.append(path)
    return sorted(templates, key=lambda p: p.name)


def resolve_template_path(template_path):
    if template_path.suffix.lower() == ".doc":
        converted_template = convert_doc_to_docx(template_path)
        if converted_template:
            return converted_template
        sibling_docx = template_path.with_suffix(".docx")
        if sibling_docx.exists():
            return sibling_docx
        return None
    return template_path


def allowed_file(filename, allowed_exts):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_exts


def normalize_text(text):
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def is_section_heading(text):
    if not text:
        return False
    text = text.strip()
    if len(text) > 20 and text.upper() == text and any(c.isalpha() for c in text):
        return True
    headings = [
        "abstract", "introduction", "related work", "methods", "methodology", "results", "discussion", "conclusion", "references", "acknowledgment",
        "摘要", "引言", "相关工作", "方法", "方法论", "结果", "讨论", "结论", "参考文献", "致谢", "致谢"
    ]
    lowered = text.lower()
    return any(lowered.startswith(h) for h in headings)


def paragraph_is_heading(paragraph):
    style_name = ""
    try:
        style_name = paragraph.style.name or ""
    except Exception:
        style_name = ""
    if style_name and "heading" in style_name.lower():
        return True
    # Detect centered + bold paragraphs as headings (common in submitted manuscripts)
    try:
        align = paragraph.alignment
        is_center = align == WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception:
        is_center = False

    try:
        has_bold = any((r.bold is True) for r in paragraph.runs if r.text and r.text.strip())
    except Exception:
        has_bold = False

    if is_center and has_bold:
        return True

    return is_section_heading(paragraph.text)


def extract_text_from_plain_text(content):
    if isinstance(content, bytes):
        content = content.decode("utf-8", errors="ignore")
    paragraphs = []
    for block in re.split(r"\n\s*\n", content):
        text = normalize_text(block)
        if not text:
            continue
        paragraph_type = "heading" if is_section_heading(text) else "normal"
        paragraphs.append((text, paragraph_type))
    return paragraphs


def extract_text_from_doc(file_path):
    paragraphs = []
    raw_parts = []
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.gencache.EnsureDispatch("Word.Application")
        try:
            word.Visible = False
        except Exception:
            pass
        try:
            word.DisplayAlerts = 0
        except Exception:
            pass
        doc = word.Documents.Open(
            str(file_path),
            True,
            True,
            False,
            "",
            "",
            True,
            "",
            "",
            0,
        )
        try:
            for para in doc.Paragraphs:
                text = normalize_text((para.Range.Text or "").strip())
                if not text:
                    continue
                paragraph_type = "heading" if is_section_heading(text) else "normal"
                paragraphs.append((text, paragraph_type))
                raw_parts.append(text)
        finally:
            try:
                doc.Close(False)
            except Exception:
                pass
            try:
                word.Quit()
            except Exception:
                pass
    except Exception as exc:
        print("extract_text_from_doc error", repr(exc))

    raw_text = "\n\n".join(raw_parts)
    if paragraphs:
        return paragraphs, raw_text
    return [], ""


def extract_text_from_docx(file_path):
    def append_paragraph(text, paragraphs, raw_parts, paragraph_obj=None):
        normalized = normalize_text(text)
        if not normalized:
            return
        if paragraph_obj is not None:
            try:
                paragraph_type = "heading" if paragraph_is_heading(paragraph_obj) else "normal"
            except Exception:
                paragraph_type = "heading" if is_section_heading(normalized) else "normal"
        else:
            paragraph_type = "heading" if is_section_heading(normalized) else "normal"
        paragraphs.append((normalized, paragraph_type))
        raw_parts.append(normalized)

    paragraphs = []
    raw_parts = []
    # Primary extraction using python-docx (paragraphs + table cells)
    try:
        document = Document(file_path)
        for p in document.paragraphs:
            append_paragraph(p.text, paragraphs, raw_parts, paragraph_obj=p)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        append_paragraph(p.text, paragraphs, raw_parts, paragraph_obj=p)
    except Exception:
        paragraphs = []
        raw_parts = []

    raw_text = "\n\n".join(raw_parts)

    # Fallback 1: docx2txt (often extracts textboxes and other parts)
    if (not paragraphs or len(raw_text.strip()) < 10) and docx2txt is not None:
        try:
            raw_text = docx2txt.process(str(file_path))
            paragraphs = extract_text_from_plain_text(raw_text)
        except Exception:
            paragraphs = []

    # Fallback 2: mammoth (HTML/plain text conversion)
    if (not paragraphs or len(raw_text.strip()) < 10) and mammoth is not None:
        try:
            res = mammoth.extract_raw_text(str(file_path))
            raw_text = res.value
            paragraphs = extract_text_from_plain_text(raw_text)
        except Exception:
            paragraphs = []

    # Fallback 3: raw XML parse from document.xml (captures text in many places)
    if not paragraphs:
        try:
            texts = []
            with zipfile.ZipFile(str(file_path)) as z:
                names = [n for n in z.namelist() if n.startswith("word/") and n.endswith(".xml")]
                for name in names:
                    try:
                        data = z.read(name)
                        tree = ET.fromstring(data)
                        for node in tree.iter():
                            if node.tag.endswith('}t') and node.text:
                                texts.append(node.text)
                    except Exception:
                        continue
            raw_text = "\n\n".join(texts)
            paragraphs = extract_text_from_plain_text(raw_text)
        except Exception:
            paragraphs = []

    return paragraphs, raw_text


def extract_docx_elements(file_path):
    """Extract paragraphs, raw_text, tables (as nested lists) and image paths from a docx file."""
    paragraphs, raw_text = extract_text_from_docx(file_path)
    tables_list = []
    images = []
    try:
        doc = Document(file_path)
        for table in doc.tables:
            rows = []
            for r in table.rows:
                cells = []
                for c in r.cells:
                    cells.append(normalize_text(c.text))
                rows.append(cells)
            tables_list.append(rows)
    except Exception:
        tables_list = []

    # extract images from word/media
    try:
        with zipfile.ZipFile(str(file_path)) as z:
            media = [n for n in z.namelist() if n.startswith('word/media/')]
            if media:
                img_dir = UPLOAD_FOLDER / 'extracted_media'
                img_dir.mkdir(parents=True, exist_ok=True)
                for name in media:
                    data = z.read(name)
                    target = img_dir / Path(name).name
                    with open(target, 'wb') as out:
                        out.write(data)
                    images.append(target)
    except Exception:
        images = []

    return paragraphs, raw_text, tables_list, images


def extract_text_from_txt(content):
    if isinstance(content, bytes):
        content = content.decode("utf-8", errors="ignore")
    return extract_text_from_plain_text(content)


def clear_document_body(document):
    body = document._element.body
    for element in list(body):
        # keep section properties (sectPr) so document.sections remains available
        try:
            tag = element.tag
        except Exception:
            tag = ''
        if tag.endswith('}sectPr'):
            continue
        body.remove(element)


def find_template_style_by_names(document, candidates):
    for name in candidates:
        try:
            _ = document.styles[name]
            return name
        except Exception:
            continue
    return None


def find_template_style(document, target):
    if target == "heading":
        candidates = [
            "H1_No Space", "H1", "Heading 1", "Heading1", "paper_heading_1", "Title1", "제목-1", "Heading(even)", "Heading(odd)",
            "Heading 2", "Heading 3", "Heading 4", "Heading 5",
        ]
        name = find_template_style_by_names(document, candidates)
        if name:
            return name
        candidates = [s.name for s in document.styles if "heading" in s.name.lower()]
        if candidates:
            return candidates[0]
    if target == "heading2":
        candidates = [
            "H2", "H2_No Space", "Heading 2", "Heading2", "paper_heading_2", "Heading(odd)", "heading2", "Heading 1",
        ]
        name = find_template_style_by_names(document, candidates)
        if name:
            return name
        return find_template_style(document, "heading")
    if target == "abstract":
        return find_template_style_by_names(document, ["Abstract", "abstract", "Abstract/Keyword", "Abstract내용-1"])
    if target == "keywords":
        return find_template_style_by_names(document, ["IT", "keywords", "Keyword", "Keywords", "Abstract/Keyword", "@키워드"])
    if target == "normal":
        return find_template_style_by_names(document, [
            "PARA", "본문-2", "Body Text", "Normal (Web)", "Normal", "Paper Title", "Text", "본문내용", "Body Text 3", "본문-1",
        ]) or "Normal"
    return None


def is_subsection_heading(text):
    if not text:
        return False
    return bool(re.match(r"^\d+(\.\d+)+", text.strip())) or bool(re.match(r"^[A-Z]\.", text.strip()))


def is_abstract_heading(text):
    lowered = text.strip().lower()
    return lowered.startswith("abstract") or lowered.startswith("摘要") or lowered.startswith("abstract:")


def is_keywords_heading(text):
    lowered = text.strip().lower()
    return lowered.startswith("keywords") or lowered.startswith("keyword") or lowered.startswith("关键词")


def build_formatted_document(article_paragraphs, template_path):
    document = Document(template_path)
    heading_style = find_template_style(document, "heading") or "Heading 1"
    heading2_style = find_template_style(document, "heading2") or heading_style
    abstract_style = find_template_style(document, "abstract")
    keywords_style = find_template_style(document, "keywords")
    normal_style = find_template_style(document, "normal") or "Normal"
    clear_document_body(document)

    for para_text, para_type in article_paragraphs:
        style = normal_style
        if para_type == "heading":
            if is_abstract_heading(para_text) and abstract_style:
                style = abstract_style
            elif is_keywords_heading(para_text) and keywords_style:
                style = keywords_style
            elif is_subsection_heading(para_text):
                style = heading2_style
            else:
                style = heading_style
        else:
            if is_abstract_heading(para_text) and abstract_style:
                style = abstract_style
            elif is_keywords_heading(para_text) and keywords_style:
                style = keywords_style
            else:
                style = normal_style
        document.add_paragraph(para_text, style=style)

    return document


def append_tables_and_images(document, tables_list, images):
    # Append tables (simple text copy) then images at end
    for tbl in tables_list:
        if not tbl:
            continue
        rows = len(tbl)
        cols = max(len(r) for r in tbl)
        table = document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, r in enumerate(tbl):
            for j, cell_text in enumerate(r):
                try:
                    table.rows[i].cells[j].text = cell_text or ''
                except Exception:
                    continue

    for img in images:
        try:
            document.add_picture(str(img), width=Inches(5))
        except Exception:
            try:
                document.add_picture(str(img))
            except Exception:
                continue


def convert_doc_to_docx(doc_path):
    output_path = doc_path.with_suffix(".docx")

    # 优先使用 Windows COM 自动化将 .doc 转换为 .docx
    try:
        import win32com.client
        import shutil
        import tempfile
        import pythoncom

        if output_path.exists():
            output_path.unlink()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_dir_path = Path(tmp_dir)
            temp_doc = tmp_dir_path / doc_path.name
            temp_out = tmp_dir_path / (doc_path.stem + ".docx")
            shutil.copy(str(doc_path), str(temp_doc))
            pythoncom.CoInitialize()
            word = win32com.client.gencache.EnsureDispatch("Word.Application")
            try:
                word.Visible = False
            except Exception:
                pass
            try:
                word.DisplayAlerts = 0
            except Exception:
                pass
            wd_format = 12
            doc = word.Documents.Open(
                str(temp_doc),
                True,
                True,
                False,
                "",
                "",
                True,
                "",
                "",
                0,
            )
            try:
                save_method = getattr(doc, "SaveAs2", None) or getattr(doc, "SaveAs", None)
                if save_method is None:
                    raise AttributeError("Word document object has no SaveAs method")
                save_method(str(temp_out.resolve()), FileFormat=wd_format)
            finally:
                try:
                    doc.Close(False)
                except Exception:
                    pass
                try:
                    word.Quit()
                except Exception:
                    pass
            if temp_out.exists():
                shutil.move(str(temp_out), str(output_path))
        if output_path.exists():
            return output_path
    except Exception as e:
        try:
            word.Quit()
        except Exception:
            pass
        print('convert_doc_to_docx error', repr(e))
        # continue to fallback options without raising

    # 回退使用 pypandoc，如果环境中安装了 Pandoc
    try:
        import pypandoc
        pypandoc.convert_file(str(doc_path), "docx", outputfile=str(output_path))
        if output_path.exists():
            return output_path
    except Exception:
        pass

    return None


@app.route("/", methods=["GET", "POST"])
def index():
    builtin_templates = get_builtin_templates()
    builtin_template_names = [path.name for path in builtin_templates]

    download_url = None
    download_name = None
    if request.method == "POST":
        article_file = request.files.get("article_file")
        template_file = request.files.get("template_file")
        selected_template_name = request.form.get("builtin_template")
        ai_endpoint = request.form.get("ai_endpoint", "").strip()
        ai_api_key = request.form.get("ai_api_key", "").strip()
        ai_model = request.form.get("ai_model", "").strip()

        if not article_file or article_file.filename == "":
            flash("请上传文章文件。")
            return redirect(request.url)

        article_filename = secure_filename(article_file.filename)
        if not allowed_file(article_filename, ALLOWED_ARTICLE_EXTENSIONS):
            flash("文章文件只支持 .doc 或 .docx 或 .txt 格式。")
            return redirect(request.url)

        article_path = UPLOAD_FOLDER / article_filename
        article_file.save(article_path)

        if template_file and template_file.filename != "":
            template_filename = secure_filename(template_file.filename)
            if not allowed_file(template_filename, ALLOWED_TEMPLATE_EXTENSIONS):
                flash("自定义模板文件只支持 .doc 或 .docx 格式。")
                return redirect(request.url)
            template_path = UPLOAD_FOLDER / template_filename
            template_file.save(template_path)
            template_path = resolve_template_path(template_path)
            if not template_path:
                flash("自定义模板为 .doc 格式，当前环境暂时无法自动转换，请先用 Word 将其保存为 .docx。")
                return redirect(request.url)
        else:
            if not selected_template_name:
                flash("请上传模板文件或选择一个内置模板。")
                return redirect(request.url)
            template_path = BUILTIN_TEMPLATE_DIR / selected_template_name
            if not template_path.exists():
                flash("所选内置模板不存在，请刷新页面后重试。")
                return redirect(request.url)
            template_path = resolve_template_path(template_path)
            if not template_path:
                flash("所选内置模板为 .doc 格式，当前环境暂时无法自动转换，请先用 Word 将其保存为 .docx。")
                return redirect(request.url)

        article_paragraphs = []
        raw_text = ""
        if article_filename.lower().endswith(".txt"):
            with open(article_path, "rb") as text_stream:
                raw_text = text_stream.read().decode("utf-8", errors="ignore")
            article_paragraphs = extract_text_from_txt(raw_text)
        elif article_filename.lower().endswith(".doc"):
            converted_article = convert_doc_to_docx(article_path)
            if converted_article:
                article_paragraphs, raw_text = extract_text_from_docx(converted_article)
            else:
                article_paragraphs, raw_text = extract_text_from_doc(article_path)
                if not article_paragraphs:
                    flash("文章为 .doc 格式，当前环境无法读取或转换，请先在 Word 中保存为 .docx。")
                    return redirect(request.url)
        else:
            article_paragraphs, raw_text = extract_text_from_docx(article_path)

        if article_paragraphs and not any(pt == "heading" for _, pt in article_paragraphs):
            ai_sections = parse_sections_with_ai(
                raw_text,
                endpoint=ai_endpoint or None,
                api_key=ai_api_key or None,
                model=ai_model or None,
            )
            if ai_sections:
                new_paragraphs = []
                for section in ai_sections:
                    heading = section.get("heading", "").strip()
                    content = section.get("content", "").strip()
                    if heading:
                        new_paragraphs.append((heading, "heading"))
                    if content:
                        for content_para, _ in extract_text_from_plain_text(content):
                            new_paragraphs.append((content_para, "normal"))
                if new_paragraphs:
                    article_paragraphs = new_paragraphs

        if not article_paragraphs:
            flash("未能从文章文件中提取有效内容，请检查文件格式。")
            return redirect(request.url)

        # build document from template and paragraphs
        formatted_document = build_formatted_document(article_paragraphs, template_path)

        # try to extract tables/images from the original article and append them
        try:
            # use converted_article if we converted earlier
            source_doc_path = None
            if article_filename.lower().endswith('.doc'):
                source_doc_path = converted_article if 'converted_article' in locals() else None
            else:
                source_doc_path = article_path
            if source_doc_path:
                _, _, tables_list, images = extract_docx_elements(source_doc_path)
                append_tables_and_images(formatted_document, tables_list, images)
        except Exception:
            pass

        output_filename = f"{Path(article_filename).stem}_formatted_{int(time.time())}.docx"
        output_path = FORMATTED_FOLDER / output_filename
        formatted_document.save(str(output_path))

        download_url = url_for("download_file", filename=output_filename)
        download_name = f"{Path(article_filename).stem}_formatted.docx"

        return render_template(
            "index.html",
            builtin_templates=builtin_template_names,
            download_url=download_url,
            download_name=download_name,
        )

    return render_template(
        "index.html",
        builtin_templates=builtin_template_names,
        download_url=download_url,
        download_name=download_name,
    )


@app.route('/download/<path:filename>')
def download_file(filename):
    safe_path = FORMATTED_FOLDER / filename
    try:
        if not safe_path.exists() or safe_path.resolve().parent != FORMATTED_FOLDER.resolve():
            flash("下载文件不存在。")
            return redirect(url_for("index"))
    except Exception:
        flash("下载文件不存在。")
        return redirect(url_for("index"))
    return send_from_directory(str(FORMATTED_FOLDER), filename, as_attachment=True)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
